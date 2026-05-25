import io
import unicodedata
from docx import Document
from docx.shared import Pt
from fpdf import FPDF


def exportar_markdown(texto: str) -> bytes:
    return texto.encode("utf-8")


def exportar_docx(texto: str) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for linha in texto.split("\n"):
        linha = linha.strip()
        if not linha:
            doc.add_paragraph("")
        elif linha.startswith("# "):
            doc.add_heading(linha[2:], level=1)
        elif linha.startswith("## "):
            doc.add_heading(linha[3:], level=2)
        elif linha.startswith("### "):
            doc.add_heading(linha[4:], level=3)
        elif linha.startswith(("- ", "* ")):
            doc.add_paragraph(linha[2:], style="List Bullet")
        else:
            doc.add_paragraph(linha.replace("**", ""))
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _texto_pdf(texto: str) -> str:
    trocas = {
        "•": "-",
        "–": "-",
        "—": "-",
        "“": '"',
        "”": '"',
        "‘": "'",
        "’": "'",
    }
    for antigo, novo in trocas.items():
        texto = texto.replace(antigo, novo)
    texto = unicodedata.normalize("NFKD", texto)
    return texto.encode("latin-1", errors="ignore").decode("latin-1")


def _multi_cell(pdf: FPDF, altura: int, texto: str) -> None:
    pdf.multi_cell(0, altura, _texto_pdf(texto))
    pdf.set_x(pdf.l_margin)


def exportar_pdf(texto: str) -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    for linha in texto.split("\n"):
        linha = linha.strip()
        if not linha:
            pdf.ln(4)
        elif linha.startswith("# "):
            pdf.set_font("Helvetica", "B", 16)
            pdf.set_text_color(27, 42, 74)
            _multi_cell(pdf, 10, linha[2:])
        elif linha.startswith("## "):
            pdf.set_font("Helvetica", "B", 13)
            pdf.set_text_color(14, 111, 86)
            _multi_cell(pdf, 8, linha[3:])
        elif linha.startswith("### "):
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(50, 50, 50)
            _multi_cell(pdf, 7, linha[4:])
        elif linha.startswith(("- ", "* ")):
            pdf.set_font("Helvetica", "", 11)
            pdf.set_text_color(30, 30, 30)
            _multi_cell(pdf, 7, f"- {linha[2:]}")
        else:
            pdf.set_font("Helvetica", "", 11)
            pdf.set_text_color(30, 30, 30)
            _multi_cell(pdf, 7, linha.replace("**", ""))
    return bytes(pdf.output())
