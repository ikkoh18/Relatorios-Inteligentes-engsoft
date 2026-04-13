import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF


def exportar_markdown(texto: str) -> bytes:
    return texto.encode("utf-8")


def exportar_docx(texto: str) -> bytes:
    doc = Document()

    # Estilos básicos
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for linha in texto.split("\n"):
        linha = linha.strip()
        if not linha:
            doc.add_paragraph("")
            continue

        if linha.startswith("# "):
            p = doc.add_heading(linha[2:], level=1)
        elif linha.startswith("## "):
            p = doc.add_heading(linha[3:], level=2)
        elif linha.startswith("### "):
            p = doc.add_heading(linha[4:], level=3)
        elif linha.startswith("- ") or linha.startswith("* "):
            doc.add_paragraph(linha[2:], style="List Bullet")
        elif linha.startswith("**") and linha.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(linha.strip("**"))
            run.bold = True
        else:
            doc.add_paragraph(linha)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def exportar_pdf(texto: str) -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    for linha in texto.split("\n"):
        linha = linha.strip()
        if not linha:
            pdf.ln(4)
            continue

        if linha.startswith("# "):
            pdf.set_font("Helvetica", "B", 16)
            pdf.set_text_color(27, 42, 74)
            pdf.multi_cell(0, 10, linha[2:])
            pdf.ln(2)
        elif linha.startswith("## "):
            pdf.set_font("Helvetica", "B", 13)
            pdf.set_text_color(14, 111, 86)
            pdf.multi_cell(0, 8, linha[3:])
            pdf.ln(1)
        elif linha.startswith("### "):
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(50, 50, 50)
            pdf.multi_cell(0, 7, linha[4:])
        elif linha.startswith("- ") or linha.startswith("* "):
            pdf.set_font("Helvetica", "", 11)
            pdf.set_text_color(30, 30, 30)
            pdf.multi_cell(0, 7, f"  • {linha[2:]}")
        else:
            pdf.set_font("Helvetica", "", 11)
            pdf.set_text_color(30, 30, 30)
            # Remove markdown bold **texto**
            linha_limpa = linha.replace("**", "")
            pdf.multi_cell(0, 7, linha_limpa)

    return bytes(pdf.output())
